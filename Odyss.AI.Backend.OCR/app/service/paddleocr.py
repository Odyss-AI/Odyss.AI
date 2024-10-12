from io import BytesIO
import zlib
from bson import ObjectId
from torch import device
import torch
from transformers import AutoProcessor, AutoModelForVision2Seq, StoppingCriteriaList #, TrOCRProcessor, VisionEncoderDecoderModel
import os
import PyPDF2
from pptxtopdf import convert as pptx_to_pdf
from docx2pdf import convert as docx_to_pdf
from app.user import TextChunk, Image
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image as PilImage
from app.config import Config
from .StoppingCriteraScores import StoppingCriteriaScores
import pytesseract
from difflib import SequenceMatcher


class OCRPaddle:
    def __init__(self):
        self.device = device("cuda" if torch.cuda.is_available() else "cpu")
        # Nougat Modell und Prozessor laden
        self.processor = AutoProcessor.from_pretrained("facebook/nougat-small")
        self.model = AutoModelForVision2Seq.from_pretrained("facebook/nougat-small")
        # self.processorFormula = TrOCRProcessor.from_pretrained("microsoft/trocr-small-handwritten")
        # self.modelFormula = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-small-handwritten")

    def extract_text(self, doc):
        file_extension = os.path.splitext(doc.name)[1].lower()

        if file_extension == ".pdf":
            self.process_pdf(doc)  # PDF-Verarbeitung aufrufen
        elif file_extension in [".docx", ".pptx"]:
            self.convert_docx_or_pptx_to_pdf(doc)  # Konvertiere DOCX/PPTX zu PDF
            self.process_pdf(doc)  # PDF-Verarbeitung aufrufen
        else:
            print("Unsupported file type")  # Überprüfen, ob der Dateityp nicht unterstützt wird

        return doc


# Convert
    def convert_docx_or_pptx_to_pdf(self, doc):
        try:
            if doc.name.endswith(".docx"):
                docx_to_pdf(doc.doclink)  # Convert DOCX to PDF
            elif doc.name.endswith(".pptx"):
                pptx_to_pdf(doc.doclink)  # Convert PPTX to PDF
        except Exception as e:
            raise Exception(f"Error during conversion: {e}")

        doc.name = doc.name.replace('.docx', '.pdf').replace('.pptx', '.pdf')
        return doc

    def convert_docx_to_pdf(self, doc):
        print(f"Starting DOCX to PDF conversion for: {doc.doclink}")
        
        try:
            doc_content = DocxDocument(doc.doclink)
            print(f"Number of paragraphs in DOCX: {len(doc_content.paragraphs)}")
            
            pdf_stream = BytesIO()
            pdf_canvas = canvas.Canvas(pdf_stream, pagesize=letter)

            for idx, paragraph in enumerate(doc_content.paragraphs):
                print(f"Writing Paragraph {idx + 1} to PDF: {paragraph.text}")
                pdf_canvas.drawString(100, 750, paragraph.text)
                pdf_canvas.showPage()

            pdf_canvas.save()
            pdf_stream.seek(0)

            pdf_length = pdf_stream.getbuffer().nbytes
            print(f"PDF Stream created with length: {pdf_length} bytes")

            doc.doclink = pdf_stream
            return pdf_stream

        except Exception as e:
            print(f"Error during DOCX to PDF conversion: {e}")
            raise


    def convert_pptx_to_pdf(self, doc):
        # Dateiinhalt in Bytes laden
        pptx_stream = BytesIO(doc.doclink.read())
        pdf_stream = BytesIO()  # Erstelle einen neuen BytesIO-Stream für die PDF

        try:
            # Konvertiere PPTX-Stream direkt zu PDF-Stream
            pptx_to_pdf(pptx_stream, pdf_stream)  # Anpassung hier
            pdf_stream.seek(0)  # Zurücksetzen des Streams auf den Anfang
        except Exception as e:
            raise Exception(f"Fehler bei der Konvertierung von PPTX zu PDF: {e}")

        return pdf_stream  # Gibt den PDF-Stream zurück

# handle formulas
    # def extract_formulas_from_image(self, image_path):
    #     processorFormula = TrOCRProcessor.from_pretrained("microsoft/trocr-small-handwritten")
    #     modelFormula = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-small-handwritten")

    #     image = Image.open(image_path).convert("RGB")
    #     pixel_values = processorFormula(image, return_tensors="pt").pixel_values
    #     generated_ids = modelFormula.generate(pixel_values)
    #     generated_text = processorFormula.batch_decode(generated_ids, skip_special_tokens=True)[0]
        
    #     # Example regex to identify LaTeX-like formulas from the recognized text
    #     formulas = re.findall(r'\$[^\$]+\$', generated_text)
    #     return formulas
    
    # def save_formulas_to_file(self, formulas, output_path):
    #     with open(output_path, "w", encoding="utf-8") as f:
    #         for formula, page_num in formulas:
    #             f.write(f"Formula on Page {page_num}: {formula}\n")

    # def extract_text_formulas_from_pdf(self, file_path):
    #     formulas = []
    #     with pdfplumber.open(file_path) as pdf:
    #         for page_num, page in enumerate(pdf.pages):
    #             text = page.extract_text()
    #             matches = re.findall(r'\$[^\$]+\$', text)  # Extract LaTeX-like formulas
    #             formulas.extend([(match, page_num + 1) for match in matches])
    #     return formulas


# extraction
    def extract_text_from_pdf(self, pdf_stream):
        full_text = ""
        page_texts = []  # Liste, um Text pro Seite zu speichern
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    full_text += f"{page_text.strip()}\n"
                    page_texts.append((page_text.strip(), page_num + 1))  # Seitenzahl hinzufügen
                else:
                    full_text += f"No text detected on page {page_num + 1}\n"
                    page_texts.append(("", page_num + 1))  # Leeren Text hinzufügen
            
            if not full_text.strip():  # If no text was found
                full_text = ""  # Return empty string for image-only documents

        except Exception as e:
            print(f"Error extracting text from PDF: {e}")

        return page_texts  # Gib die Liste von Seiten zurück


    def extract_images_from_pdf(self, pdf_stream, doc):
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            for page_num, page in enumerate(pdf_reader.pages):
                resources = page.get("/Resources").get_object()
                xobjects = resources.get("/XObject")
                
                if xobjects:
                    xobjects = xobjects.get_object()
                    image_counter = 1  # Image counter für jede Seite zurücksetzen
                    
                    for obj in xobjects:
                        xobject = xobjects[obj].get_object()
                        
                        if xobject["/Subtype"] == "/Image":
                            try:
                                width = xobject["/Width"]
                                height = xobject["/Height"]

                                # Daten extrahieren
                                data = xobject._data
                                file_extension = "png"  # Standard PNG verwenden

                                # Überprüfe auf vorhandene Filter und dekodiere entsprechend
                                if "/Filter" in xobject:
                                    if xobject["/Filter"] == "/FlateDecode":
                                        try:
                                            data = zlib.decompress(data)  # Dekomprimiere die FlateDecode-Daten
                                            # Erstelle ein Bild aus den dekomprimierten Daten
                                            img = PilImage.frombytes("RGB", (width, height), data)
                                        except Exception as e:
                                            print(f"Fehler beim Dekomprimieren von Bild {image_counter} auf Seite {page_num + 1}: {e}")
                                            continue
                                    elif xobject["/Filter"] == "/DCTDecode":
                                        file_extension = "jpg"  # JPEG benötigt keine Decodierung
                                    elif xobject["/Filter"] == "/JPXDecode":
                                        file_extension = "jp2"  # JPEG2000
                                    else:
                                        print(f"Unbekannter Filter {xobject['/Filter']} für Bild {image_counter} auf Seite {page_num + 1}")
                                        continue  # Überspringe unbekannte Filter

                                # Speicherpfad für das Bild
                                img_save_path = os.path.join(Config.LOCAL_DOC_PATH, f"extracted_image_{page_num+1}_{image_counter}.{file_extension}")

                                # Speichere das Bild basierend auf dem Filtertyp
                                img.save(img_save_path)
                                print(f"Bild {image_counter} auf Seite {page_num + 1} erfolgreich gespeichert als {img_save_path}.")

                                # OCR auf dem Bild ausführen
                                print(f"Starte OCR für Bild {image_counter} auf Seite {page_num + 1}...")
                                img_text = self.ocr_image(img_save_path)  # OCR-Funktion

                                # if img_text.strip() == "":
                                #     print(f"OCR-Ergebnis leer für Bild {image_counter} auf Seite {page_num + 1}.")
                                # else:
                                #     print(f"OCR-Ergebnis für Bild {image_counter} auf Seite {page_num + 1}: {img_text}")

                                # Erstelle ein Image-Objekt
                                image_obj = Image(
                                id=str(ObjectId()),
                                link=img_save_path,
                                page=page_num + 1,
                                type=file_extension,
                                imgtext=img_text if isinstance(img_text, str) else "",
                                llm_output=""
                            )
                                doc.imgList.append(image_obj)

                                image_counter += 1

                            except Exception as e:
                                print(f"Fehler beim Verarbeiten des Bildes für Objekt {obj} auf Seite {page_num + 1}: {e}")

        except Exception as e:
            print(f"Fehler beim Extrahieren der Bilder aus dem PDF: {e}")

        print(f"Image extraction complete. Images found: {len(doc.imgList)}")


    def split_text_into_chunks(self, full_text, doc, page_num):
            # Split based on double newlines to capture paragraphs or sections
            chunks = full_text.split('\n\n')
            for chunk in chunks:
                if chunk.strip():
                    text_chunk = TextChunk(id=str(ObjectId()), text=chunk.strip(), page=page_num)
                    doc.textList.append(text_chunk)

    def ocr_nougat(self, image_stream):
        try:
            print(f"Starte Bildverarbeitung für OCR...")

            # Lade das Bild aus dem Stream
            image = PilImage.open(image_stream).convert("RGB")

            # Bild für das Modell vorbereiten
            pixel_values = self.processor(images=image, return_tensors="pt").pixel_values

            # Textextraktion durchführen mit benutzerdefinierter StoppingCriteria
            outputs = self.model.generate(
                pixel_values.to(self.device),
                min_length=1,
                max_length=3584,
                bad_words_ids=[[self.processor.tokenizer.unk_token_id]],  # Filter für unbekannte Tokens
                return_dict_in_generate=True,
                output_scores=True,
                stopping_criteria=StoppingCriteriaList([StoppingCriteriaScores()])  
            )

            # Überprüfen, ob Ausgabe generiert wurde
            if not outputs or len(outputs[0]) == 0:
                return ""  # Leeren String zurückgeben, wenn nichts erkannt wurde

            # Ergebnis dekodieren
            generated_text = self.processor.batch_decode(outputs[0], skip_special_tokens=True)[0]

            # Postprocess the generation (optional, je nach Anwendungsfall)
            generated_text = self.processor.post_process_generation(generated_text, fix_markdown=False)

            # Gebe den erkannten Text zurück, wenn vorhanden, ansonsten leeren String
            if generated_text.strip():
                print(f"Texterkennung abgeschlossen. Erkannt: {generated_text}")
                return generated_text
            else:
                print(f"Kein Text erkannt.")
                return ""

        except Exception as e:
            print(f"Fehler bei der OCR für Bild: {e}")
            return ""  # Stelle sicher, dass "" zurückgegeben wird, falls ein Fehler auftritt
        
        

    def ocr_image(self, image_stream):
        try:
            print(f"Starte Tesseract-Bildverarbeitung für OCR...")

            # Lade das Bild aus dem Stream
            image = PilImage.open(image_stream)
            print(f"Bild erfolgreich geladen.")

            # Textextraktion mit Tesseract durchführen
            tesseract_text = pytesseract.image_to_string(image)
            print(f"Tesseract erkannter Text: {tesseract_text}")

            return tesseract_text.strip() if tesseract_text.strip() else ""

        except Exception as e:
            print(f"Fehler bei Tesseract OCR für Bild: {e}")
            return ""



    def compare_ocr_results(self, text1, text2):
        ratio = SequenceMatcher(None, text1, text2).ratio()
        print(f"Ähnlichkeit der OCR-Ergebnisse: {ratio*100:.2f}%")
        return ratio

    # def ocr_image(self, image_stream):
    #     try:
    #         print(f"Starte Bildverarbeitung für OCR...")

    #         # Nougat OCR
    #         nougat_text = self.ocr_nougat(image_stream)

    #         # Tesseract OCR
    #         tesseract_text = self.ocr_tesseract(image_stream)

    #         # Vergleich der Ergebnisse
    #         similarity = self.compare_ocr_results(nougat_text, tesseract_text)

    #         return {"nougat": nougat_text, "tesseract": tesseract_text, "similarity": similarity}

    #     except Exception as e:
    #         print(f"Fehler bei der OCR für Bild: {e}")
    #         return {"nougat": "", "tesseract": "", "similarity": 0}

    # def ocr_image(self, image_stream):

            # Nur Nougat OCR
            # self.ocr_nougat(image_stream)

            #
            # self.ocr_tesseract(image_stream)

            # Extrahiere nur den Text
            # imgtext = nougat_result if isinstance(nougat_result, str) else nougat_result.get("imgtext", "")

            # Speichere nur den erkannten Text im imgtext-Feld
            # return {"imgtext": imgtext}


# Sudo main
    def process_pdf(self, doc):
        try:
            # Attempt to extract text from PDF
            page_texts = self.extract_text_from_pdf(doc.doclink)

            print("Splitting extracted text into chunks...")
            for text, page_num in page_texts:
                self.split_text_into_chunks(text, doc, page_num)  # page_num wird jetzt korrekt übergeben
            print(f"Text chunks created: {len(doc.textList)}")

            # Always attempt image extraction regardless of text outcome
            print("Attempting to extract images from PDF...")
            self.extract_images_from_pdf(doc.doclink, doc)

        except Exception as e:
            print(f"Error during PDF processing: {e}")

