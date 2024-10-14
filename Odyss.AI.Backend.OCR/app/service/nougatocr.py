from io import BytesIO
from bson import ObjectId
from torch import device
import torch
from transformers import AutoProcessor, AutoModelForVision2Seq
import os
from docx2pdf import convert as docx_to_pdf
from pptxtopdf import convert as pptx_to_pdf
from app.user import TextChunk
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image as PilImage
from app.config import Config

class OCRNougat:
    def __init__(self):
        self.device = device("cuda" if torch.cuda.is_available() else "cpu")
        # Nougat Modell und Prozessor laden
        self.processor = AutoProcessor.from_pretrained("facebook/nougat-small")
        self.model = AutoModelForVision2Seq.from_pretrained("facebook/nougat-small")

    def extract_text(self, doc):
        file_extension = os.path.splitext(doc.name)[1].lower()

        if file_extension == ".pdf":
            self.process_pdf(doc)  # PDF-Verarbeitung aufrufen
        elif file_extension in [".docx", ".pptx"]:
            self.convert_docx_or_pptx_to_pdf(doc)  # Konvertiere DOCX/PPTX zu PDF
            self.process_pdf(doc)  # PDF-Verarbeitung aufrufen
        else:
            print("Unsupported file type")  # Überprüfen, ob der Dateityp nicht unterstützt wird

        return doc

    def convert_docx_or_pptx_to_pdf(self, doc):
        try:
            if doc.name.endswith(".docx"):
                docx_to_pdf(doc.doclink)  # Convert DOCX to PDF
            elif doc.name.endswith(".pptx"):
                pptx_to_pdf(doc.doclink)  # Convert PPTX to PDF
        except Exception as e:
            raise Exception(f"Error during conversion: {e}")

        doc.name = doc.name.replace('.docx', '.pdf').replace('.pptx', '.pdf')
        return doc

    def convert_docx_to_pdf(self, doc):
        print(f"Starting DOCX to PDF conversion for: {doc.doclink}")
        
        try:
            doc_content = DocxDocument(doc.doclink)
            print(f"Number of paragraphs in DOCX: {len(doc_content.paragraphs)}")
            
            pdf_stream = BytesIO()
            pdf_canvas = canvas.Canvas(pdf_stream, pagesize=letter)

            for idx, paragraph in enumerate(doc_content.paragraphs):
                print(f"Writing Paragraph {idx + 1} to PDF: {paragraph.text}")
                pdf_canvas.drawString(100, 750, paragraph.text)
                pdf_canvas.showPage()

            pdf_canvas.save()
            pdf_stream.seek(0)

            pdf_length = pdf_stream.getbuffer().nbytes
            print(f"PDF Stream created with length: {pdf_length} bytes")

            doc.doclink = pdf_stream
            return pdf_stream

        except Exception as e:
            print(f"Error during DOCX to PDF conversion: {e}")
            raise

    def process_pdf(self, doc):
        try:
            print(f"Processing PDF: {doc.name}")
            pdf_stream = doc.doclink  # PDF als BytesIO oder Stream
            self.run_ocr(pdf_stream, doc)
        except Exception as e:
            print(f"Error processing PDF: {e}")

    def run_ocr(self, pdf_stream, doc):
        try:
            # Bereite das Modell und den Prozessor für die OCR-Verarbeitung des PDFs vor
            print(f"Running OCR on the PDF...")

            # Lade das PDF in Nougat zur Texterkennung
            pixel_values = self.processor(pdf_stream, return_tensors="pt").pixel_values
            outputs = self.model.generate(pixel_values.to(self.device))

            # Texterkennung dekodieren
            generated_text = self.processor.batch_decode(outputs, skip_special_tokens=True)[0]

            print(f"OCR text: {generated_text}")

            # Teile den Text in Chunks
            self.split_text_into_chunks(generated_text, doc)

        except Exception as e:
            print(f"Error during OCR processing: {e}")

    def split_text_into_chunks(self, full_text, doc):
        chunks = full_text.split('\n\n')  # Aufteilen in Absätze
        for chunk in chunks:
            if chunk.strip():
                text_chunk = TextChunk(id=str(ObjectId()), text=chunk.strip(), page=1)  # Beispiel: Seite 1
                doc.textList.append(text_chunk)
