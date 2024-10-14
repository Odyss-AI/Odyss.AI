from io import BytesIO
import zlib
from bson import ObjectId
import os
import PyPDF2
from pptxtopdf import convert as pptx_to_pdf
from docx2pdf import convert as docx_to_pdf
from app.user import TextChunk, Image
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image as PilImage
from app.config import Config
import pytesseract


class OCRTesseract:
    def extract_text(self, doc):
        file_extension = os.path.splitext(doc.name)[1].lower()

        if file_extension == ".pdf":
            self.process_pdf(doc)  # PDF-Verarbeitung aufrufen
        elif file_extension in [".docx", ".pptx"]:
            self.convert_docx_or_pptx_to_pdf(doc)  # Konvertiere DOCX/PPTX zu PDF
            self.process_pdf(doc)  # PDF-Verarbeitung aufrufen
        else:
            print("Unsupported file type")  # Überprüfen, ob der Dateityp nicht unterstützt wird

        return doc

    def convert_docx_or_pptx_to_pdf(self, doc):
        try:
            if doc.name.endswith(".docx"):
                docx_to_pdf(doc.doclink)  # Convert DOCX to PDF
            elif doc.name.endswith(".pptx"):
                pptx_to_pdf(doc.doclink)  # Convert PPTX to PDF
        except Exception as e:
            raise Exception(f"Error during conversion: {e}")

        doc.name = doc.name.replace('.docx', '.pdf').replace('.pptx', '.pdf')
        return doc

    def convert_docx_to_pdf(self, doc):
        print(f"Starting DOCX to PDF conversion for: {doc.doclink}")
        
        try:
            doc_content = DocxDocument(doc.doclink)
            print(f"Number of paragraphs in DOCX: {len(doc_content.paragraphs)}")
            
            pdf_stream = BytesIO()
            pdf_canvas = canvas.Canvas(pdf_stream, pagesize=letter)

            for idx, paragraph in enumerate(doc_content.paragraphs):
                print(f"Writing Paragraph {idx + 1} to PDF: {paragraph.text}")
                pdf_canvas.drawString(100, 750, paragraph.text)
                pdf_canvas.showPage()

            pdf_canvas.save()
            pdf_stream.seek(0)

            pdf_length = pdf_stream.getbuffer().nbytes
            print(f"PDF Stream created with length: {pdf_length} bytes")

            doc.doclink = pdf_stream
            return pdf_stream

        except Exception as e:
            print(f"Error during DOCX to PDF conversion: {e}")
            raise

    def extract_text_from_pdf(self, pdf_stream):
        full_text = ""
        page_texts = []  # Liste, um Text pro Seite zu speichern
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    full_text += f"{page_text.strip()}\n"
                    page_texts.append((page_text.strip(), page_num + 1))  # Seitenzahl hinzufügen
                else:
                    full_text += f"No text detected on page {page_num + 1}\n"
                    page_texts.append(("", page_num + 1))  # Leeren Text hinzufügen
            
            if not full_text.strip():  # If no text was found
                full_text = ""  # Return empty string for image-only documents

        except Exception as e:
            print(f"Error extracting text from PDF: {e}")

        return page_texts  # Gib die Liste von Seiten zurück

    def extract_images_from_pdf(self, pdf_stream, doc):
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            for page_num, page in enumerate(pdf_reader.pages):
                resources = page.get("/Resources").get_object()
                xobjects = resources.get("/XObject")
                
                if xobjects:
                    xobjects = xobjects.get_object()
                    image_counter = 1  # Image counter für jede Seite zurücksetzen
                    
                    for obj in xobjects:
                        xobject = xobjects[obj].get_object()
                        
                        if xobject["/Subtype"] == "/Image":
                            try:
                                width = xobject["/Width"]
                                height = xobject["/Height"]

                                # Daten extrahieren
                                data = xobject._data
                                file_extension = "png"  # Standard PNG verwenden

                                # Überprüfe auf vorhandene Filter und dekodiere entsprechend
                                if "/Filter" in xobject:
                                    if xobject["/Filter"] == "/FlateDecode":
                                        try:
                                            data = zlib.decompress(data)  # Dekomprimiere die FlateDecode-Daten
                                            img = PilImage.frombytes("RGB", (width, height), data)
                                        except Exception as e:
                                            print(f"Fehler beim Dekomprimieren von Bild {image_counter} auf Seite {page_num + 1}: {e}")
                                            continue
                                    elif xobject["/Filter"] == "/DCTDecode":
                                        file_extension = "jpg"
                                    elif xobject["/Filter"] == "/JPXDecode":
                                        file_extension = "jp2"
                                    else:
                                        print(f"Unbekannter Filter {xobject['/Filter']} für Bild {image_counter} auf Seite {page_num + 1}")
                                        continue

                                # Speicherpfad für das Bild
                                img_save_path = os.path.join(Config.LOCAL_DOC_PATH, f"extracted_image_{page_num+1}_{image_counter}.{file_extension}")

                                # Speichere das Bild basierend auf dem Filtertyp
                                img.save(img_save_path)
                                print(f"Bild {image_counter} auf Seite {page_num + 1} erfolgreich gespeichert als {img_save_path}.")

                                # OCR auf dem Bild ausführen
                                print(f"Starte OCR für Bild {image_counter} auf Seite {page_num + 1}...")
                                img_text = self.ocr_image(img_save_path)  # Tesseract OCR-Funktion

                                # Erstelle ein Image-Objekt
                                image_obj = Image(
                                    id=str(ObjectId()),
                                    link=img_save_path,
                                    page=page_num + 1,
                                    type=file_extension,
                                    imgtext=img_text if isinstance(img_text, str) else "",
                                    llm_output=""
                                )
                                doc.imgList.append(image_obj)

                                image_counter += 1

                            except Exception as e:
                                print(f"Fehler beim Verarbeiten des Bildes für Objekt {obj} auf Seite {page_num + 1}: {e}")

        except Exception as e:
            print(f"Fehler beim Extrahieren der Bilder aus dem PDF: {e}")

        print(f"Image extraction complete. Images found: {len(doc.imgList)}")

    def split_text_into_chunks(self, full_text, doc, page_num):
        chunks = full_text.split('\n\n')
        for chunk in chunks:
            if chunk.strip():
                text_chunk = TextChunk(id=str(ObjectId()), text=chunk.strip(), page=page_num)
                doc.textList.append(text_chunk)

    def ocr_image(self, image_stream):
        try:
            print(f"Starte Tesseract-Bildverarbeitung für OCR...")

            # Lade das Bild aus dem Stream
            image = PilImage.open(image_stream)
            print(f"Bild erfolgreich geladen.")

            # Textextraktion mit Tesseract durchführen
            tesseract_text = pytesseract.image_to_string(image)

            if tesseract_text.strip():
                print(f"OCR erfolgreich. Erkannt: {tesseract_text}")
                return tesseract_text
            else:
                print(f"Kein Text erkannt.")
                return ""

        except Exception as e:
            print(f"Fehler bei der Tesseract OCR: {e}")
            return ""  # Stelle sicher, dass "" zurückgegeben wird, falls ein Fehler auftritt
